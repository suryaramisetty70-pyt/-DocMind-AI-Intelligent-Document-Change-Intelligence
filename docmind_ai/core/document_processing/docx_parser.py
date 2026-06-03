"""
DocMind AI - DOCX Parser
Handles Word document parsing
"""

from pathlib import Path
from typing import List, Dict, Any
from .base_parser import (
    BaseDocumentParser,
    DocumentContent,
    DocumentMetadata,
    DocumentElement,
    DocumentParserFactory
)

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    docx = None
    DOCX_AVAILABLE = False


class DOCXParser(BaseDocumentParser):
    """Parser for Word documents (DOCX)"""
    
    SUPPORTED_EXTENSIONS = ["docx"]
    
    def __init__(self):
        super().__init__()
    
    def parse(self, file_path: Path) -> DocumentContent:
        """Parse DOCX file and extract structured content"""
        self.validate_file(file_path)
        
        metadata = self.extract_metadata(file_path)
        text_content = self.extract_text(file_path)
        elements = self._extract_elements(file_path)
        structure = self._analyze_structure(elements)
        
        metadata.page_count = 1
        
        return DocumentContent(
            metadata=metadata,
            text_content=text_content,
            elements=elements,
            tables=[],
            images=[],
            structure=structure
        )
    
    def extract_text(self, file_path: Path) -> str:
        """Extract text from DOCX"""
        if not DOCX_AVAILABLE:
            return "DOCX parsing not available. Install python-docx."
        
        try:
            doc = docx.Document(file_path)
            paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
            return "\n\n".join(paragraphs)
        except Exception as e:
            return f"Error parsing DOCX: {str(e)}"
    
    def _extract_elements(self, file_path: Path) -> List[DocumentElement]:
        """Extract structured elements from DOCX"""
        elements = []
        
        if not DOCX_AVAILABLE:
            return elements
        
        try:
            doc = docx.Document(file_path)
            
            for i, para in enumerate(doc.paragraphs, 1):
                if para.text.strip():
                    element = DocumentElement(
                        element_type="text",
                        content=para.text,
                        page_number=1,
                        position=(0, i * 20, 100, (i + 1) * 20),
                        confidence=1.0
                    )
                    elements.append(element)
            
            for table in doc.tables:
                element = DocumentElement(
                    element_type="table",
                    content=str([[cell.text for cell in row.cells] for row in table.rows]),
                    page_number=1
                )
                elements.append(element)
        
        except Exception:
            pass
        
        return elements
    
    def _analyze_structure(self, elements: List[DocumentElement]) -> Dict[str, Any]:
        """Analyze document structure"""
        return {
            "sections": [{"title": f"Section {i+1}", "level": 1} for i in range(len(elements) // 5 + 1)],
            "total_elements": len(elements)
        }


DocumentParserFactory.register_parser("docx", DOCXParser())
